"""Convert an existing resume into a structured :class:`Profile`.

Supports plain text formats (``.md``, ``.txt``), Word ``.docx`` files, and any
PDF (LinkedIn's "Save to PDF" export is just a normal PDF). For ``.doc`` we
ask the user to convert to ``.docx`` because the legacy binary format would
require a fragile native dependency.

The CLI also lets users paste resume text directly (``--paste`` /
``init`` interactive fallback), in which case we skip the file-reading
step. Either path ends in :func:`extract_profile_from_text`, which asks
the configured LLM to fill out the :class:`Profile` schema via
``with_structured_output``. There is no markdown fallback any more — a
working LLM is required at ``init`` time.
"""

from __future__ import annotations

from collections.abc import Callable
from pathlib import Path
from typing import TYPE_CHECKING

from langchain_core.messages import HumanMessage, SystemMessage

from jobapply.config import AppConfig, get_api_key
from jobapply.profile import Profile

if TYPE_CHECKING:  # pragma: no cover - import for type hints only
    from langchain_core.language_models.chat_models import BaseChatModel

SUPPORTED_SUFFIXES: tuple[str, ...] = (".md", ".txt", ".docx", ".pdf")


class ResumeImportError(RuntimeError):
    """Raised when a resume can't be read or converted into a Profile."""


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency declared in pyproject
        raise ResumeImportError(
            "python-docx is required to read .docx files. "
            "Run `pip install python-docx` or reinstall the package.",
        ) from exc

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency declared in pyproject
        raise ResumeImportError(
            "pypdf is required to read PDF files. "
            "Run `pip install pypdf` or reinstall the package.",
        ) from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # pragma: no cover - rare malformed PDFs
            text = ""
        text = text.strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages).strip()


def extract_text_from_resume(path: Path) -> str:
    """Return raw text for ``path`` based on its extension."""
    if not path.is_file():
        raise ResumeImportError(f"No such file: {path}")
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _read_text_file(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".doc":
        raise ResumeImportError(
            "Legacy .doc isn't supported. Open the file and re-save as .docx, then retry.",
        )
    raise ResumeImportError(
        f"Unsupported resume format '{suffix}'. Use one of: {', '.join(SUPPORTED_SUFFIXES)}.",
    )


_SYSTEM_PROMPT = (
    "You convert a candidate's existing resume into a STRUCTURED JSON "
    "profile. The schema you produce will be the source of truth for "
    "downstream resume tailoring, so accuracy beats embellishment.\n\n"
    "RULES (must follow):\n"
    "- Preserve every fact verbatim: do not invent employers, dates, "
    "metrics, schools, GPAs, links, or skills. If the source doesn't "
    "mention something, leave that field empty rather than guessing.\n"
    "- Drop boilerplate like recruiter notes, page numbers, and watermarks.\n"
    "- Copy URLs and email addresses character-for-character; do NOT "
    "shorten them.\n"
    "- Classify each link by host into the dedicated fields: a github.com "
    "URL goes into `github`, linkedin.com into `linkedin`, medium.com into "
    "`medium`, twitter.com or x.com into `twitter`, and a personal site "
    "into `portfolio`. Anything else goes into `other_links` with a short "
    "human-readable label. Bare usernames are also acceptable for the "
    "named fields (e.g. `linkedin: 'jane-doe'`).\n"
    "- `skills` is a FLAT list of distinct skills. Do not nest, "
    "categorize, or duplicate; if the source groups them ('Languages: "
    "Python, Go'), unwrap the group and add each skill on its own. "
    "Preserve parenthesized aliases such as 'Model Context Protocol "
    "(MCP)'.\n"
    "- For each `experience` entry: split the date range into "
    "`start_date` and `end_date` (use 'Present' for current roles). "
    "Convert the bullet points into the `bullets` list, one bullet per "
    "entry, lightly cleaned (no leading dashes).\n"
    "- For each `education` entry: split the date range, capture GPA in "
    "the `gpa` field with its scale (e.g. '9.6/10' or '3.85/4.0') and "
    "split coursework into a list. Use `honors` for awards / thesis text.\n"
    "- For each `project`: pull a one-line summary into `description`, "
    "the canonical link into `url`, and the tech stack into `tech` "
    "(separate from `bullets`).\n"
    "- The `summary` field is the candidate's existing professional "
    "summary verbatim (or the most relevant equivalent). If the resume "
    "has no summary, leave it empty.\n"
)


def llm_extract_profile(llm: BaseChatModel, raw_text: str) -> Profile:
    """Ask the LLM to populate a :class:`Profile` from ``raw_text``.

    Uses ``with_structured_output(Profile)`` so providers that support
    function/tool calling return a directly-validated Pydantic instance.
    Wrapped exceptions surface as :class:`ResumeImportError` for the CLI.
    """
    structured = llm.with_structured_output(Profile)
    user = HumanMessage(content=f"RESUME SOURCE TEXT:\n\n{raw_text}")
    try:
        result = structured.invoke([SystemMessage(content=_SYSTEM_PROMPT), user])
    except Exception as exc:  # noqa: BLE001 - re-raise as a single CLI error type
        raise ResumeImportError(
            f"LLM structured-output extraction failed "
            f"({type(exc).__name__}: {exc}). Re-run `jobapply init` once your "
            "provider is reachable.",
        ) from exc
    if not isinstance(result, Profile):
        raise ResumeImportError(
            f"LLM returned {type(result).__name__}, expected a Profile. "
            "The provider's structured-output support may be broken — try a "
            "different model or provider.",
        )
    return result


def _ensure_llm_factory(
    cfg: AppConfig,
    llm_factory: Callable[[AppConfig, str, str], BaseChatModel] | None,
) -> Callable[[AppConfig, str, str], BaseChatModel]:
    """Return ``llm_factory`` or build the default one lazily."""
    if llm_factory is not None:
        return llm_factory
    from jobapply.llm import create_chat_model

    def _default_factory(c: AppConfig, p: str, m: str) -> BaseChatModel:
        return create_chat_model(p, m, cfg=c)

    return _default_factory


def extract_profile_from_text(
    raw_text: str,
    cfg: AppConfig,
    *,
    llm_factory: Callable[[AppConfig, str, str], BaseChatModel] | None = None,
) -> Profile:
    """Ask the configured provider to turn ``raw_text`` into a :class:`Profile`.

    ``cfg`` must point at a provider with a working API key (or be
    ``ollama``). Otherwise we raise immediately with a message telling the
    user how to configure one.
    """
    if not raw_text or not raw_text.strip():
        raise ResumeImportError(
            "Resume text is empty. Paste your resume content or pass a "
            "non-empty file via --resume.",
        )

    provider = cfg.provider
    if provider != "ollama" and not get_api_key(cfg, provider):
        raise ResumeImportError(
            f"No API key configured for provider '{provider}'. Run "
            "`jobapply config` (or set the matching env var) before "
            "running `jobapply init` so the resume can be parsed into "
            "profile.json.",
        )

    factory = _ensure_llm_factory(cfg, llm_factory)
    model = cfg.resolved_model(provider)
    llm = factory(cfg, provider, model)
    return llm_extract_profile(llm, raw_text)


def extract_profile_from_resume(
    resume_path: Path,
    cfg: AppConfig,
    *,
    llm_factory: Callable[[AppConfig, str, str], BaseChatModel] | None = None,
) -> Profile:
    """Read ``resume_path`` and convert it to a :class:`Profile` via the LLM."""
    raw_text = extract_text_from_resume(resume_path)
    if not raw_text.strip():
        raise ResumeImportError(
            f"Could not extract any text from {resume_path}. "
            "If it's a scanned PDF, try a Markdown or DOCX export instead.",
        )
    return extract_profile_from_text(raw_text, cfg, llm_factory=llm_factory)
